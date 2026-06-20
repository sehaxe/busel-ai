"""
🛰️ busel — Any-to-Token Encoders (vocab=326)

Sovereign byte-level encoders for images, video, audio, PDF, and docx.

Busel is byte-level: the same `BitLinear_a4_8` processes every modality.
The encoders below turn files into a stream of integer tokens that fits
into the dynamically-computed vocab (see `multimodal.special_tokens`).
Decoders reverse the process for inspection.

CRITICAL — representation:
    The encoders return `list[int]` (NOT `bytes`). The values 0-255 are
    real bytes; 256+ are RESERVED TOKEN INDICES in the model's vocab
    (256-258 legacy + 259+ plug-in specials from `special_tokens`).
    Python's `bytes` type cannot represent values >= 256, so the multimodal
    stream is an `int` stream that the data collate function converts to
    `int32` tensors.

Multimodal layout (v5.4, plug-in token IDs from `special_tokens`):
    IMAGE:   [MOD_IMAGE] [3072 raw RGB bytes @ 32x32] [MEDIA_END]
    VIDEO:   [MOD_VIDEO] [4-byte frame_count LE] (frame_0 ... frame_N) [MEDIA_END]
             each frame = [3072 raw RGB bytes]
    AUDIO:   [MOD_AUDIO] [4-byte sample_rate LE] [4-byte n_samples LE]
             [2-byte sample_width] [little-endian int16 PCM] [MEDIA_END]
    PDF:     [MOD_PDF]   [Docling-converted UTF-8 text] [MEDIA_END]
    DOCX:    [MOD_DOCX]  [python-docx plain-text UTF-8] [MEDIA_END]
    TEXT:    [MOD_TEXT]  [UTF-8 bytes, no MEDIA_END — unbounded]

Legacy layouts (v5.0-5.3, still decoded for backward compat):
    [MEDIA_START] [payload] [MEDIA_END]  — no MOD_* prefix

All `MOD_*` IDs and `MEDIA_*` IDs come from
`multimodal.special_tokens` — no hardcoded numbers in this file.
"""
from __future__ import annotations

import os
import struct
import wave
from io import BytesIO
from typing import Tuple

from busel_registry import register
from multimodal.special_tokens import (
    MEDIA_START,
    MEDIA_END,
    DOC_SEP,
    is_enabled,
    get_special_token,
)

# Backward-compat alias — older code (v5.0-5.3) imports `IMAGE_MARKER` to mean
# the generic "media start" token. The plug-in split makes the modality explicit
# (MOD_IMAGE, MOD_VIDEO, ...) but we keep the legacy name for the same value.
IMAGE_MARKER = MEDIA_START

IMAGE_W = 32
IMAGE_H = 32
IMAGE_BYTES = IMAGE_W * IMAGE_H * 3

try:
    import cv2
    import numpy as np
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import imageio.v3 as iio
    HAS_IMAGEIO = True
except ImportError:
    HAS_IMAGEIO = False

try:
    import soundfile as sf
    HAS_SOUNDFILE = True
except ImportError:
    HAS_SOUNDFILE = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from docling.document_converter import DocumentConverter
    HAS_DOCLING = True
except ImportError:
    HAS_DOCLING = False


def _resolve_modality_marker(modality: str) -> int:
    """Resolve a MOD_* token id, falling back to legacy MEDIA_START if disabled.

    If the modality's special token has been disabled at runtime, we emit the
    legacy `MEDIA_START` (256) so the stream is still valid for the model.
    The downstream `MEDIA_END` is always emitted (legacy, never disabled).
    """
    name_map = {
        "image": "mod_image",
        "video": "mod_video",
        "audio": "mod_audio",
        "pdf":   "mod_pdf",
        "docx":  "mod_docx",
        "text":  "mod_text",
    }
    if modality not in name_map:
        raise ValueError(f"unknown modality: {modality!r}")
    name = name_map[modality]
    if is_enabled(name):
        return get_special_token(name).id
    return MEDIA_START


@register("encoder", "image")
class ImageEncoder:
    """Encode/decode RGB images @ 32×32 = 3072 token payload.

    Fast path uses OpenCV (cv2.imdecode + cv2.resize INTER_AREA + cv2.cvtColor),
    ~3× faster than PIL on realistic 1024² images. PIL is a fallback when cv2
    is unavailable.

    Layout: [MOD_IMAGE] [3072 raw RGB bytes] [MEDIA_END]
    """

    name = "image"
    extensions = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff")

    def __init__(self, size: Tuple[int, int] = (IMAGE_W, IMAGE_H)):
        if not (HAS_CV2 or HAS_PIL):
            raise ImportError("opencv-python-headless OR Pillow required: uv add opencv-python-headless pillow")
        self.size = size

    def encode_file(self, path: str) -> list:
        if HAS_CV2:
            arr = cv2.imread(path, cv2.IMREAD_COLOR)
            if arr is None:
                arr = cv2.imdecode(np.fromfile(path, dtype=np.uint8), cv2.IMREAD_COLOR)
            if arr is None:
                if HAS_PIL:
                    return self._encode_pil_path(path)
                raise ValueError(f"cv2 failed to decode {path}")
            arr = cv2.resize(arr, self.size, interpolation=cv2.INTER_AREA)
            arr = cv2.cvtColor(arr, cv2.COLOR_BGR2RGB)
        else:
            return self._encode_pil_path(path)
        raw = arr.tobytes()
        assert len(raw) == IMAGE_BYTES, f"image bytes {len(raw)} != {IMAGE_BYTES}"
        return [_resolve_modality_marker("image")] + list(raw) + [MEDIA_END]

    def _encode_pil_path(self, path: str) -> list:
        with Image.open(path) as im:
            im = im.convert("RGB").resize(self.size)
            raw = im.tobytes()
        assert len(raw) == IMAGE_BYTES
        return [_resolve_modality_marker("image")] + list(raw) + [MEDIA_END]

    def encode(self, image) -> list:
        if HAS_CV2 and not (HAS_PIL and isinstance(image, Image.Image)):
            if isinstance(image, (bytes, bytearray)):
                arr = cv2.imdecode(np.frombuffer(bytes(image), dtype=np.uint8), cv2.IMREAD_COLOR)
            else:
                arr = np.asarray(image)
                if arr.ndim == 3 and arr.shape[2] == 3:
                    arr = cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)
            if arr is None:
                raise ValueError("cv2.imdecode failed on image bytes")
            arr = cv2.resize(arr, self.size, interpolation=cv2.INTER_AREA)
            arr = cv2.cvtColor(arr, cv2.COLOR_BGR2RGB)
        else:
            im = image.convert("RGB").resize(self.size)
            arr = np.asarray(im)
        raw = arr.tobytes()
        assert len(raw) == IMAGE_BYTES, f"image bytes {len(raw)} != {IMAGE_BYTES}"
        return [_resolve_modality_marker("image")] + list(raw) + [MEDIA_END]

    def decode(self, tokens: list) -> "Image.Image":
        if not tokens:
            raise ValueError("empty token stream")
        if tokens[0] not in (get_special_token("mod_image").id, MEDIA_START):
            raise ValueError(f"tokens missing image marker (got {tokens[0]})")
        try:
            end = tokens.index(MEDIA_END, 1)
        except ValueError:
            raise ValueError("tokens missing media-end marker")
        payload = tokens[1:end]
        if len(payload) != IMAGE_BYTES:
            raise ValueError(f"image payload {len(payload)} != {IMAGE_BYTES}")
        return Image.frombytes("RGB", self.size, bytes(payload))


@register("encoder", "video")
class VideoEncoder:
    """Encode videos as a sequence of 32×32 RGB frame payloads.

    Fast path uses OpenCV VideoCapture with CAP_PROP_FRAME_COUNT (single
    metadata call) + cap.grab() for seek-skipping. imageio fallback
    iterates the video twice (frame count + frames) and is ~5-10× slower.

    Layout: [MOD_VIDEO] [4-byte frame_count LE] (frame_0 ... frame_N) [MEDIA_END]
    """

    name = "video"
    extensions = (".mp4", ".mov", ".avi", ".mkv", ".webm")

    def __init__(self, size: Tuple[int, int] = (IMAGE_W, IMAGE_H), max_frames: int = 8):
        if not (HAS_CV2 or HAS_IMAGEIO):
            raise ImportError("opencv-python-headless OR imageio required")
        self.size = size
        self.max_frames = max_frames

    def _resize_frame(self, frame: "np.ndarray") -> list:
        if HAS_CV2:
            arr = cv2.resize(frame, self.size, interpolation=cv2.INTER_AREA)
            arr = cv2.cvtColor(arr, cv2.COLOR_BGR2RGB)
        else:
            from PIL import Image as _Im
            pil = _Im.fromarray(frame).convert("RGB").resize(self.size)
            arr = np.asarray(pil)
        return list(arr.tobytes())

    def encode_file(self, path: str) -> list:
        marker = _resolve_modality_marker("video")
        if HAS_CV2:
            cap = cv2.VideoCapture(path)
            if not cap.isOpened():
                cap.release()
                if HAS_IMAGEIO:
                    return self._encode_imageio(path, marker)
                raise ValueError(f"cv2.VideoCapture failed to open {path}")
            n_total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            if n_total <= 0:
                cap.release()
                return [marker, MEDIA_END]
            step = max(1, n_total // self.max_frames)
            out = [marker]
            out += list(struct.pack("<I", self.max_frames))
            n_written = 0
            for idx in range(n_total):
                if idx % step != 0:
                    cap.grab()
                    continue
                if n_written >= self.max_frames:
                    break
                ok, frame = cap.read()
                if not ok:
                    break
                out += self._resize_frame(frame)
                n_written += 1
            cap.release()
            out.append(MEDIA_END)
            return out
        return self._encode_imageio(path, marker)

    def _encode_imageio(self, path: str, marker: int) -> list:
        n_total = 0
        for _ in iio.imiter(path):
            n_total += 1
        if n_total == 0:
            return [marker, MEDIA_END]
        step = max(1, n_total // self.max_frames)
        out = [marker]
        out += list(struct.pack("<I", self.max_frames))
        n_written = 0
        for idx, frame in enumerate(iio.imiter(path)):
            if idx % step != 0:
                continue
            if n_written >= self.max_frames:
                break
            out += self._resize_frame(frame)
            n_written += 1
        out.append(MEDIA_END)
        return out

    def decode(self, tokens: list) -> list:
        if not tokens:
            raise ValueError("empty token stream")
        if tokens[0] not in (get_special_token("mod_video").id, MEDIA_START):
            raise ValueError(f"tokens missing video marker (got {tokens[0]})")
        n = struct.unpack("<I", bytes(tokens[1:5]))[0]
        cursor = 5
        frames = []
        for _ in range(n):
            payload = tokens[cursor:cursor + IMAGE_BYTES]
            if len(payload) != IMAGE_BYTES:
                break
            frames += payload
            cursor += IMAGE_BYTES
        return frames


@register("encoder", "audio")
class AudioEncoder:
    """Encode WAV / FLAC / OGG audio as 16-bit PCM token stream.

    Layout: [MOD_AUDIO] [4-byte sr] [4-byte n] [2-byte sw] [int16 PCM LE] [MEDIA_END]
    """

    name = "audio"
    extensions = (".wav", ".flac", ".ogg")

    def __init__(self, max_seconds: float = 8.0, target_sr: int = 16000):
        if not HAS_SOUNDFILE:
            raise ImportError("soundfile required: uv add soundfile")
        self.max_seconds = max_seconds
        self.target_sr = target_sr

    def encode_file(self, path: str) -> list:
        data, sr = sf.read(path, dtype="float32", always_2d=False)
        if data.ndim > 1:
            data = data.mean(axis=1)
        if len(data) > int(self.max_seconds * sr):
            data = data[: int(self.max_seconds * sr)]
        pcm16 = (data * 32767.0).clip(-32768, 32767).astype("<i2").tobytes()
        marker = _resolve_modality_marker("audio")
        out = [marker]
        out += list(struct.pack("<I", sr))
        out += list(struct.pack("<I", len(pcm16) // 2))
        out += list(struct.pack("<H", 2))
        out += list(pcm16)
        out.append(MEDIA_END)
        return out

    def decode_to_wav(self, tokens: list) -> bytes:
        if not tokens:
            raise ValueError("empty token stream")
        if tokens[0] not in (get_special_token("mod_audio").id, MEDIA_START):
            raise ValueError(f"tokens missing audio marker (got {tokens[0]})")
        sr, n, sw = struct.unpack("<IIH", bytes(tokens[1:11]))
        cursor = 11
        pcm = bytes(tokens[cursor:cursor + n * sw])
        buf = BytesIO()
        with wave.open(buf, "wb") as wf:
            wf.setnchannels(1)
            wf.setsampwidth(sw)
            wf.setframerate(sr)
            wf.writeframes(pcm)
        return buf.getvalue()


@register("encoder", "pdf")
class PDFEncoder:
    """Encode PDF documents as UTF-8 text via Docling.

    Layout: [MOD_PDF] [Docling markdown UTF-8] [MEDIA_END]
    """

    name = "pdf"
    extensions = (".pdf",)

    def __init__(self):
        if not HAS_DOCLING:
            raise ImportError("docling required: uv add docling")
        self._converter = None

    def _get_converter(self):
        if self._converter is None:
            self._converter = DocumentConverter()
        return self._converter

    def encode_file(self, path: str) -> list:
        result = self._get_converter().convert(path)
        text = result.document.export_to_markdown()
        marker = _resolve_modality_marker("pdf")
        return [marker] + list(text.encode("utf-8", errors="replace")) + [MEDIA_END]


@register("encoder", "docx")
class DocxEncoder:
    """Encode .docx files as UTF-8 plain text.

    Layout: [MOD_DOCX] [python-docx plain text UTF-8] [MEDIA_END]
    """

    name = "docx"
    extensions = (".docx",)

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx required: uv add python-docx")

    def encode_file(self, path: str) -> list:
        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        marker = _resolve_modality_marker("docx")
        return [marker] + list(text.encode("utf-8", errors="replace")) + [MEDIA_END]


@register("encoder", "text")
class TextEncoder:
    """Trivial text encoder (UTF-8). Used as the default for unknown formats.

    Layout: [MOD_TEXT] [raw UTF-8 bytes]   (no MEDIA_END — unbounded)
    """

    name = "text"
    extensions = (".txt", ".md", ".py", ".json", ".jsonl", ".rs", ".cpp", ".h", ".go")

    def encode_file(self, path: str) -> list:
        with open(path, "rb") as f:
            return [_resolve_modality_marker("text")] + list(f.read())


ENCODER_REGISTRY: dict = {}


def build_encoder_for(path: str):
    """Pick the right encoder by file extension. Falls back to text."""
    ext = os.path.splitext(path)[1].lower()
    for cls in (ImageEncoder, VideoEncoder, AudioEncoder, PDFEncoder, DocxEncoder, TextEncoder):
        if ext in cls.extensions:
            try:
                return cls()
            except ImportError:
                continue
    return TextEncoder()


def auto_encode(path: str) -> list:
    """Convenience: encode any supported file → token stream."""
    return build_encoder_for(path).encode_file(path)


def list_encoders() -> list:
    """List all registered encoder names."""
    from busel_registry import list_registered
    return list_registered("encoder")
